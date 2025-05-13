import sys
import pandas as pd
import cv2
import openpyxl
import face_recognition
import numpy as np
from docx import Document
from fpdf import FPDF
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QPushButton, 
    QTableWidget, QTableWidgetItem, QTabWidget, QFileDialog, QLabel, 
    QLineEdit, QHBoxLayout, QMessageBox, QListWidget, QListWidgetItem, 
    QInputDialog, QFrame, QScrollArea, QSizePolicy
)
from PyQt6.QtGui import QPixmap, QImage, QIcon, QFont, QPalette, QColor
from PyQt6.QtCore import Qt, QTimer, QSize
from datetime import datetime, timedelta
from db_config import SessionLocal
from models import User, Attendance

class ModernButton(QPushButton):
    def __init__(self, text, icon=None):
        super().__init__(text)
        if icon:
            self.setIcon(icon)
        self.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
            QPushButton:pressed {
                background-color: #0D47A1;
            }
        """)

class ModernTable(QTableWidget):
    def __init__(self):
        super().__init__()
        self.setStyleSheet("""
            QTableWidget {
                background-color: white;
                alternate-background-color: #f5f5f5;
                border: 1px solid #ddd;
                border-radius: 5px;
                gridline-color: #ddd;
            }
            QHeaderView::section {
                background-color: #f8f9fa;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QTableWidget::item {
                padding: 5px;
            }
        """)
        self.setAlternatingRowColors(True)
        self.horizontalHeader().setStretchLastSection(True)
        self.verticalHeader().setVisible(False)

class ModernListWidget(QListWidget):
    def __init__(self):
        super().__init__()
        self.setStyleSheet("""
            QListWidget {
                background-color: white;
                border: 1px solid #ddd;
                border-radius: 5px;
                padding: 5px;
            }
            QListWidget::item {
                padding: 10px;
                border-bottom: 1px solid #eee;
            }
            QListWidget::item:selected {
                background-color: #e3f2fd;
                color: #1976D2;
            }
        """)

class FaceIDApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Face ID Attendance System")
        self.setGeometry(100, 100, 1200, 800)
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QTabWidget::pane {
                border: 1px solid #ddd;
                border-radius: 5px;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #f8f9fa;
                border: 1px solid #ddd;
                padding: 10px 20px;
                margin-right: 2px;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            }
            QTabBar::tab:selected {
                background-color: white;
                border-bottom-color: white;
            }
        """)
        
        self.db = SessionLocal()
        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)
        self.last_seen = {}
        self.last_recognized = {}
        
        self.create_database_tab()
        self.create_export_tab()
        self.create_user_management_tab()
        self.create_camera_tab()

    def create_database_tab(self):
        self.db_tab = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # Header
        header = QLabel("📊 Database Management")
        header.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: bold;
                color: #1976D2;
                margin-bottom: 20px;
            }
        """)
        layout.addWidget(header)

        # Table
        self.table = ModernTable()
        layout.addWidget(self.table)

        # Buttons
        button_layout = QHBoxLayout()
        self.clear_history_button = ModernButton("🗑 Clear History")
        self.reset_ids_button = ModernButton("🔄 Reset IDs")
        self.clear_history_button.clicked.connect(self.clear_history)
        self.reset_ids_button.clicked.connect(self.reset_ids)
        button_layout.addWidget(self.clear_history_button)
        button_layout.addWidget(self.reset_ids_button)
        layout.addLayout(button_layout)

        self.load_database()
        self.db_tab.setLayout(layout)
        self.tabs.addTab(self.db_tab, "📋 Database")

    def create_export_tab(self):
        self.export_tab = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # Header
        header = QLabel("📤 Export Data")
        header.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: bold;
                color: #1976D2;
                margin-bottom: 20px;
            }
        """)
        layout.addWidget(header)

        # Export buttons
        self.export_excel_button = ModernButton("📊 Export to Excel")
        self.export_word_button = ModernButton("📄 Export to Word")
        self.export_pdf_button = ModernButton("📑 Export to PDF")
        
        self.export_excel_button.clicked.connect(self.export_to_excel)
        self.export_word_button.clicked.connect(self.export_to_word)
        self.export_pdf_button.clicked.connect(self.export_to_pdf)
        
        layout.addWidget(self.export_excel_button)
        layout.addWidget(self.export_word_button)
        layout.addWidget(self.export_pdf_button)
        layout.addStretch()

        self.export_tab.setLayout(layout)
        self.tabs.addTab(self.export_tab, "📤 Export")

    def create_user_management_tab(self):
        self.user_tab = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # Header
        header = QLabel("👥 User Management")
        header.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: bold;
                color: #1976D2;
                margin-bottom: 20px;
            }
        """)
        layout.addWidget(header)

        # User list
        self.user_list = ModernListWidget()
        layout.addWidget(self.user_list)

        # Buttons
        button_layout = QHBoxLayout()
        self.add_user_button = ModernButton("➕ Add User")
        self.delete_user_button = ModernButton("❌ Delete User")
        self.add_user_button.clicked.connect(self.add_user)
        self.delete_user_button.clicked.connect(self.delete_user)
        button_layout.addWidget(self.add_user_button)
        button_layout.addWidget(self.delete_user_button)
        layout.addLayout(button_layout)

        self.load_users()
        self.user_tab.setLayout(layout)
        self.tabs.addTab(self.user_tab, "👥 Users")

    def create_camera_tab(self):
        self.camera_tab = QWidget()
        layout = QVBoxLayout()
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # Header
        header = QLabel("📷 Face Recognition")
        header.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: bold;
                color: #1976D2;
                margin-bottom: 20px;
            }
        """)
        layout.addWidget(header)

        # Camera display
        self.camera_label = QLabel()
        self.camera_label.setMinimumSize(640, 480)
        self.camera_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.camera_label.setStyleSheet("""
            QLabel {
                background-color: #000;
                border-radius: 10px;
            }
        """)
        layout.addWidget(self.camera_label)

        # Camera controls
        button_layout = QHBoxLayout()
        self.start_camera_button = ModernButton("📷 Start Camera")
        self.stop_camera_button = ModernButton("🛑 Stop Camera")
        self.start_camera_button.clicked.connect(self.start_camera)
        self.stop_camera_button.clicked.connect(self.stop_camera)
        button_layout.addWidget(self.start_camera_button)
        button_layout.addWidget(self.stop_camera_button)
        layout.addLayout(button_layout)

        self.camera_tab.setLayout(layout)
        self.tabs.addTab(self.camera_tab, "📷 Camera")

    def load_database(self):
        records = self.db.query(Attendance).all()
        self.table.setRowCount(len(records))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["ID", "Имя", "Дата", "Время"])

        for row_idx, record in enumerate(records):
            self.table.setItem(row_idx, 0, QTableWidgetItem(str(record.id)))
            self.table.setItem(row_idx, 1, QTableWidgetItem(record.name))
            self.table.setItem(row_idx, 2, QTableWidgetItem(record.date))
            self.table.setItem(row_idx, 3, QTableWidgetItem(record.time))

    def clear_history(self):
        self.db.query(Attendance).delete()
        self.db.commit()
        self.load_database()
        QMessageBox.information(self, "Очистка", "История успешно очищена!")

    def reset_ids(self):
        records = self.db.query(Attendance).order_by(Attendance.id).all()
        self.db.query(Attendance).delete()
        for idx, record in enumerate(records, start=1):
            new_record = Attendance(
                id=idx,
                name=record.name,
                date=record.date,
                time=record.time
            )
            self.db.add(new_record)
        self.db.commit()
        self.load_database()
        QMessageBox.information(self, "Обновление ID", "ID успешно пересчитаны!")

    def export_to_excel(self):
        records = self.db.query(Attendance).all()
        df = pd.DataFrame([(r.id, r.name, r.date, r.time) for r in records],
                         columns=['ID', 'Имя', 'Дата', 'Время'])
        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить как", "", "Excel Files (*.xlsx)")
        if file_path:
            df.to_excel(file_path, index=False)

    def export_to_word(self):
        records = self.db.query(Attendance).all()
        doc = Document()
        doc.add_heading("История посещений", level=1)
        for record in records:
            doc.add_paragraph(f"ID: {record.id}, Имя: {record.name}, Дата: {record.date}, Время: {record.time}")
        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить как", "", "Word Files (*.docx)")
        if file_path:
            doc.save(file_path)

    def export_to_pdf(self):
        records = self.db.query(Attendance).all()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, "История посещений", ln=True, align='C')
        pdf.ln(10)
        for record in records:
            pdf.cell(200, 10, f"ID: {record.id}, Имя: {record.name}, Дата: {record.date}, Время: {record.time}", ln=True)
        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить как", "", "PDF Files (*.pdf)")
        if file_path:
            pdf.output(file_path)

    def load_users(self):
        self.user_list.clear()
        users = self.db.query(User).all()
        for user in users:
            item = QListWidgetItem(user.name)
            if user.photo:
                pixmap = QPixmap()
                pixmap.loadFromData(user.photo)
                icon = QIcon(pixmap.scaled(50, 50, Qt.AspectRatioMode.KeepAspectRatio))
                item.setIcon(icon)
            self.user_list.addItem(item)

    def add_user(self):
        while True:
            name, ok = QInputDialog.getText(self, "Добавление пользователя", "Введите имя:")
            if not ok or not name:
                return

            if any(char.isdigit() for char in name):
                QMessageBox.warning(self, "Ошибка", "Имя не должно содержать цифр!")
                continue

            break  # Если имя прошло проверку, выходим из цикла

        file_path, _ = QFileDialog.getOpenFileName(self, "Выберите фото", "", "Images (*.png *.jpg *.jpeg)")
        if file_path:
            self.save_user(name, file_path)

    def save_user(self, name, file_path):
        with open(file_path, 'rb') as f:
            photo_data = f.read()
        new_user = User(name=name, photo=photo_data)
        self.db.add(new_user)
        self.db.commit()
        self.load_users()

    def delete_user(self):
        selected_item = self.user_list.currentItem()
        if selected_item:
            name = selected_item.text()
            self.db.query(User).filter(User.name == name).delete()
            self.db.commit()
            self.load_users()

    def start_camera(self):
        # Try to open the built-in camera (usually index 0)
        self.cap = cv2.VideoCapture(0, cv2.CAP_AVFOUNDATION)  # Use AVFoundation for macOS
        
        if not self.cap.isOpened():
            QMessageBox.warning(self, "Error", "Could not open the built-in camera!")
            return

        # Set camera properties for better performance
        self.cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
        self.cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
        self.cap.set(cv2.CAP_PROP_FPS, 30)
        self.cap.set(cv2.CAP_PROP_BUFFERSIZE, 1)  # Minimize buffer size
        self.cap.set(cv2.CAP_PROP_FOURCC, cv2.VideoWriter_fourcc('M', 'J', 'P', 'G'))  # Use MJPEG for better performance

        # Initialize frame processing variables
        self.frame_count = 0
        self.last_processed_time = datetime.now()
        self.processing_interval = timedelta(milliseconds=100)  # Process every 100ms
        self.known_face_encodings = []
        self.known_face_names = []
        self.load_known_faces()  # Pre-load face encodings

        # Create timer with shorter interval for smoother video
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_camera)
        self.timer.start(16)  # ~60 FPS (1000ms/16ms)

    def load_known_faces(self):
        """Pre-load all known face encodings for faster recognition"""
        users = self.db.query(User).all()
        self.known_face_encodings = []
        self.known_face_names = []

        for user in users:
            try:
                np_array = np.frombuffer(user.photo, dtype=np.uint8)
                image = cv2.imdecode(np_array, cv2.IMREAD_COLOR)
                encodings = face_recognition.face_encodings(image)
                if encodings:
                    self.known_face_encodings.append(encodings[0])
                    self.known_face_names.append(user.name)
            except Exception as e:
                print(f"Error loading image for {user.name}: {e}")

    def update_camera(self):
        if not hasattr(self, 'cap') or not self.cap.isOpened():
            return

        # Skip frames if we're falling behind
        for _ in range(2):
            self.cap.grab()

        ret, frame = self.cap.read()
        if not ret:
            return

        # Resize frame for better performance
        frame = cv2.resize(frame, (640, 480))
        frame = cv2.flip(frame, 1)
        
        # Convert to RGB for face recognition
        rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        
        # Process face recognition at fixed intervals
        current_time = datetime.now()
        if current_time - self.last_processed_time >= self.processing_interval:
            self.last_processed_time = current_time
            self.process_face_recognition(rgb_frame)

        # Convert to QImage for display
        h, w, ch = rgb_frame.shape
        bytes_per_line = ch * w
        qt_img = QImage(rgb_frame.data, w, h, bytes_per_line, QImage.Format.Format_RGB888)

        # Scale image to fit label while preserving aspect ratio
        pixmap = QPixmap.fromImage(qt_img)
        scaled_pixmap = pixmap.scaled(
            self.camera_label.width(),
            self.camera_label.height(),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.FastTransformation
        )

        self.camera_label.setPixmap(scaled_pixmap)

    def process_face_recognition(self, frame):
        """Process face recognition separately to avoid blocking the UI"""
        face_locations = face_recognition.face_locations(frame, model="hog")  # Use HOG for better performance
        if not face_locations:
            return

        face_encodings = face_recognition.face_encodings(frame, face_locations)
        
        for face_encoding in face_encodings:
            matches = face_recognition.compare_faces(self.known_face_encodings, face_encoding, tolerance=0.6)
            name = "Неизвестный"

            if True in matches:
                first_match_index = matches.index(True)
                name = self.known_face_names[first_match_index]
                print(f"Recognized: {name}")
                
                now = datetime.now()
                if name not in self.last_seen or now - self.last_seen[name] > timedelta(minutes=1):
                    self.last_seen[name] = now
                    self.save_attendance(name)
                    self.load_database()
            else:
                self.show_unrecognized_face_warning()

    def show_unrecognized_face_warning(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Icon.Warning)
        msg.setWindowTitle("Face Not Recognized")
        msg.setText("Face not found in database. Please register!")
        msg.setStyleSheet("""
            QMessageBox {
                background-color: white;
            }
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 5px 15px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        msg.exec()

    def save_attendance(self, name):
        now = datetime.now()
        date = now.strftime("%Y-%m-%d")
        time = now.strftime("%H:%M:%S")
        new_attendance = Attendance(name=name, date=date, time=time)
        self.db.add(new_attendance)
        self.db.commit()
        print(f"{name} записан в базу на {date} {time}")

    def stop_camera(self):
        if hasattr(self, 'timer'):
            self.timer.stop()
        if hasattr(self, 'cap') and self.cap.isOpened():
            self.cap.release()
        self.camera_label.clear()
        # Clear face recognition data
        self.known_face_encodings = []
        self.known_face_names = []
        if hasattr(self, 'frame_count'):
            delattr(self, 'frame_count')
        if hasattr(self, 'last_processed_time'):
            delattr(self, 'last_processed_time')

    def __del__(self):
        self.db.close()

# Запуск приложения
if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # Set application-wide font
    font = QFont("Inter", 10)
    app.setFont(font)
    
    window = FaceIDApp()
    window.show()
    sys.exit(app.exec())
